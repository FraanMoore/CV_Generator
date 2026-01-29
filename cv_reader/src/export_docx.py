from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Optional, cast

import docx
from docx.document import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.models import CVMaster

COLOR_TITLE = RGBColor(0x22, 0x7A, 0xAF)
COLOR_TEXT  = RGBColor(0x00, 0x00, 0x00)
COLOR_SUB_TITLE = RGBColor(0x53, 0x53, 0x53)
FONT_NAME = "Avenir Next"
FONT_NAME_DEMI_BOLD = "Avenir Next Demi Bold"
FONT_NAME_MEDIUM = "Avenir Next Medium"
FONT_SIZE_NAME = 22
FONT_SIZE_TITLES = 12
FONT_SIZE_BODY = 10
FONT_SIZE_HYPERLINK = 9

def _set_page_layout(
    doc: DocxDocument,
    *,
    left_mm: float = 10,
    right_mm: float = 15,
    top_mm: float = 15,
    bottom_mm: float = 10,
    page_w_mm: float = 210,
    page_h_mm: float = 297,
) -> None:
    for section in doc.sections:
        section.page_width = Mm(page_w_mm)
        section.page_height = Mm(page_h_mm)

        section.left_margin = Mm(left_mm)
        section.right_margin = Mm(right_mm)
        section.top_margin = Mm(top_mm)
        section.bottom_margin = Mm(bottom_mm)

# ============================================================
# HyperLinks
# ============================================================

def _add_hyperlink(
    paragraph: Paragraph,
    text: str,
    url: str,
    *,
    font_name: str = FONT_NAME,
    font_size: float = FONT_SIZE_HYPERLINK,
    color: RGBColor = COLOR_TEXT,
    underline: bool = False,
):
    """
    Adds a clickable hyperlink to a paragraph.
    """

    # Create relationship
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # Create hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")

    # Create run
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    # Font name
    if font_name:
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_pr.append(r_fonts)

    # Font size (Word uses half-points)
    if font_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size * 2)))
        r_pr.append(sz)

        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(int(font_size * 2)))
        r_pr.append(sz_cs)

    # Color
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        r_pr.append(c)

    #Not bold
    b = OxmlElement("w:b")
    b.set(qn("w:val"), "0")
    r_pr.append(b)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    r_pr.append(u)

    # Assemble
    r.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text # type: ignore[attr-defined]
    r.append(t)

    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def _set_footer_hyperlink_i18n(
    doc: DocxDocument,
    *,
    text_es: str,
    text_en: str,
    url: str,
    lang: str,
    font_name: str = FONT_NAME,
    font_size: float = FONT_SIZE_HYPERLINK,
    color: RGBColor = COLOR_SUB_TITLE,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    """
    Sets a footer that contains ONLY a hyperlink text,
    choosing Spanish or English based on `lang`.
    """
    text = text_es if lang == "es" else text_en

    for section in doc.sections:
        footer = section.footer

        for p in footer.paragraphs:
            p.clear()

        p = footer.paragraphs[0]
        p.alignment = align

        _add_hyperlink(
            p,
            text=text,
            url=url,
            font_name=font_name,
            font_size=font_size,
            color=color,
            underline=True,
        )

def _replace_placeholder_title_hyperlink(
    doc: DocxDocument,
    placeholder: str,
    *,
    text: str,
    url: str,
    icon: str = "🔗",
    font_name: str = FONT_NAME_DEMI_BOLD,
    font_size: float = FONT_SIZE_TITLES,
    color: RGBColor = COLOR_TITLE,
) -> bool:
    token = f"{{{{{placeholder}}}}}"

    for p in _iter_paragraphs(doc):
        if token not in p.text:
            continue

        _clear_paragraph(p)

        _add_hyperlink(
            p,
            text=text,
            url=url,
            font_name=font_name,
            font_size=font_size,
            color=color,
            underline=True,
        )
        r = p.add_run(f" {icon}")
        r.font.name = font_name
        r.font.size = Pt(font_size)
        r.font.color.rgb = color
        r.bold = False

        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return True

    return False

# ============================================================
# Styles
# ============================================================

def set_font(run, font_name, size):
    run.font.size = Pt(size)
    run.bold = False

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

def _resolve_style_name(doc: DocxDocument, preferred: str) -> Optional[str]:
    """
    Return the real style name in the document for `preferred`, tolerating a trailing '*'
    (Word UI may show "Bullet 2*" but the style name is usually "Bullet 2").
    """
    pref = preferred.strip().rstrip("*").strip().lower()
    for s in doc.styles:
        name = getattr(s, "name", None)
        if not name:
            continue
        if name.strip().rstrip("*").strip().lower() == pref:
            return name
    return None

def _first_existing_style(doc: DocxDocument, candidates: list[str]) -> Optional[str]:
    """Return first candidate style that exists in the document, else None."""
    existing = {s.name for s in doc.styles if getattr(s, "name", None)}
    for c in candidates:
        name = c.strip()
        if name in existing:
            return name

    for c in candidates:
        resolved = _resolve_style_name(doc, c)
        if resolved:
            return resolved
    return None

def _apply_font_to_run(
    run,
    font_name: str,
    font_size_pt: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color

def _apply_font_to_paragraph(
    p: Paragraph,
    font_name: str,
    font_size_pt: float | None = None,
    *,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    for run in p.runs:
        _apply_font_to_run(run, font_name, font_size_pt, bold, color)

# ============================================================
# Paragraph utilities (including tables)
# ============================================================

def _iter_paragraphs(doc: DocxDocument) -> Iterable[Paragraph]:
    """Iterate paragraphs in the document, including paragraphs inside tables."""
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

def _clear_paragraph(p: Paragraph) -> None:
    """Remove all runs from a paragraph without deleting the paragraph itself."""
    p_elm = p._p
    for r in list(p_elm.r_lst):
        p_elm.remove(r)

def _insert_paragraph_after(anchor: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    """
    Insert a new paragraph immediately after `anchor`.
    Keeps it simple and stable for python-docx + templates.
    """
    new_p_elm = OxmlElement("w:p")
    anchor._p.addnext(new_p_elm)  # type: ignore[attr-defined]
    new_p = cast(CT_P, new_p_elm)
    p = Paragraph(new_p, anchor._parent)  # type: ignore[arg-type, attr-defined]

    if style:
        p.style = style
    if text:
        p.add_run(text)
        _apply_font_to_paragraph(p, FONT_NAME, FONT_SIZE_BODY)
    return p

def _strip_leading_bullet_markers(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"^[-•]+\s*", "", s)
    return s.strip()

def _add_bold_paragraph_after(
    anchor: Paragraph,
    text: str,
    *,
    style: str | None,
    font_name: str,
    font_size: float,
    color: RGBColor | None,
) -> Paragraph:
    p = _insert_paragraph_after(anchor, text="", style=style)
    _clear_paragraph(p)

    r = p.add_run(text.strip())
    _apply_font_to_run(
        r,
        font_name=font_name,
        font_size_pt=font_size,
        bold=True,
        color=color,
    )

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p

def _clear_run_char_spacing(run) -> None:
    rPr = run._r.get_or_add_rPr()
    spacing = rPr.find(qn("w:spacing"))
    if spacing is not None:
        rPr.remove(spacing)

# ============================================================
# Placeholder replacement
# ============================================================

def _replace_placeholder_name(
    doc: DocxDocument,
    placeholder: str,
    text: str,
    font_name: str = FONT_NAME_DEMI_BOLD,
    font_size: float = FONT_SIZE_NAME,
    color: RGBColor | None = None,
) -> bool:
    token = f"{{{{{placeholder}}}}}"

    for p in _iter_paragraphs(doc):
        if token in p.text:
            _clear_paragraph(p)

            run = p.add_run(text)

            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = False

            if color:
                run.font.color.rgb = color

            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return True

    return False

def _replace_placeholder_title(
    doc: DocxDocument,
    placeholder: str,
    text: str,
    font_size: float = FONT_SIZE_TITLES,
    color: RGBColor | None = None,
) -> bool:
    token = f"{{{{{placeholder}}}}}"

    for p in _iter_paragraphs(doc):
        if token in p.text:
            _clear_paragraph(p)

            run = p.add_run(text)
            run.font.name = FONT_NAME_MEDIUM
            run.font.size = Pt(font_size)
            run.bold = False

            if color:
                run.font.color.rgb = color

            return True

    return False

def _replace_placeholder_contact(
    doc: DocxDocument,
    placeholder: str,
    email: str,
    location: str,
    phone: str,
    linkedin_url: str,
    github_url: str,
    color: RGBColor | None = None,
) -> bool:
    token = f"{{{{{placeholder}}}}}"

    font_name = FONT_NAME
    font_size = FONT_SIZE_HYPERLINK
    text_color = COLOR_TEXT

    for p in _iter_paragraphs(doc):
        if token not in p.text:
            continue

        _clear_paragraph(p)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        def add_text(txt: str):
            r = p.add_run(txt)
            r.font.name = font_name
            r.font.size = Pt(font_size)
            r.bold = False
            r.font.color.rgb = text_color

        add_text(email)
        add_text(" | ")
        add_text(location)
        add_text(" | ")
        add_text(phone)
        add_text(" | ")

        _add_hyperlink(
            p,
            text="LinkedIn",
            url=linkedin_url,
            font_name=font_name,
            font_size=font_size,
            color=text_color,
            underline=True,
        )

        add_text(" | ")

        _add_hyperlink(
            p,
            text="GitHub",
            url=github_url,
            font_name=font_name,
            font_size=font_size,
            color=text_color,
            underline=True,
        )

        return True

    return False

def _replace_placeholder_paragraph(doc: DocxDocument, placeholder: str, paragraph_text: str) -> bool:
    """
    Replace {{PLACEHOLDER}} with a single paragraph (no bullets).
    """
    token = f"{{{{{placeholder}}}}}"
    for p in _iter_paragraphs(doc):
        if token in p.text:
            _clear_paragraph(p)
            p.add_run(paragraph_text)
            _apply_font_to_paragraph(p, FONT_NAME, FONT_SIZE_BODY)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1
            return True
    return False

def _replace_placeholder_inline_left(
    doc: DocxDocument,
    placeholder: str,
    text: str,
    *,
    font_name: str = FONT_NAME,
    font_size: float = FONT_SIZE_BODY,
) -> bool:
    token = f"{{{{{placeholder}}}}}"
    for p in _iter_paragraphs(doc):
        if token in p.text:
            _clear_paragraph(p)

            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            r = p.add_run(text)
            _apply_font_to_run(r, font_name=font_name, font_size_pt=font_size, bold=False, color=COLOR_TEXT)
            _clear_run_char_spacing(r)

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            return True
    return False

def _replace_placeholder_experience_blocks(
    doc: DocxDocument,
    placeholder: str,
    cv: CVMaster,
    lang: str,
    bullets_per_experience: dict[str, list[str]] | None,
    job_title_style: str | None = None,
    bullet_style_preferred: str = "Bullet 2*",
    font_name: str = FONT_NAME,
    font_size: float = FONT_SIZE_BODY,
    job_title_color: RGBColor = COLOR_SUB_TITLE,
    bullet_color: RGBColor = COLOR_TEXT,
) -> bool:
    token = f"{{{{{placeholder}}}}}"

    bullet_style = (
        _resolve_style_name(doc, bullet_style_preferred)
        or _first_existing_style(doc, ["List Bullet 2", "List Bullet", "List Paragraph"])
    )

    job_style = _resolve_style_name(doc, job_title_style) if job_title_style else None

    for p in _iter_paragraphs(doc):
        if token not in p.text:
            continue

        _clear_paragraph(p)
        anchor = p
        first = True

        for exp in cv.experience:
            role = exp.role.es if lang == "es" else exp.role.en
            location = exp.location.es if lang == "es" else exp.location.en
            years = f"{exp.start_year}-{exp.end_year}"
            title_line = f"{role}, {exp.company}. {location} - {years}"

            # Job title (bold + color)
            if first:
                r = anchor.add_run(title_line)
                _apply_font_to_run(
                r,
                font_name=font_name,
                font_size_pt=font_size,
                bold=True,
                color=job_title_color,
                )

                anchor.paragraph_format.space_before = Pt(0)
                anchor.paragraph_format.space_after = Pt(0)
                anchor.paragraph_format.line_spacing = 1.0
                first = False
            else:
                anchor = _add_bold_paragraph_after(
                    anchor,
                    title_line,
                    style=job_style,
                    font_name=font_name,
                    font_size=font_size,
                    color=job_title_color,
                )

            default_bullets = exp.bullets.es if lang == "es" else exp.bullets.en
            use_bullets = (bullets_per_experience or {}).get(exp.company, default_bullets) or []

            for b in use_bullets:
                bb = _strip_leading_bullet_markers(b)
                if not bb:
                    continue

                bullet_p = _insert_paragraph_after(anchor, text="", style=bullet_style)

                prefix = ""
                if not bullet_style:
                    prefix = "• "
                    bullet_p.paragraph_format.left_indent = Pt(18)
                    bullet_p.paragraph_format.first_line_indent = Pt(-10)

                r = bullet_p.add_run(prefix + bb)
                _apply_font_to_run(
                    r,
                    font_name=font_name,
                    font_size_pt=font_size,
                    bold=False,
                    color=bullet_color,
                )

                bullet_p.paragraph_format.space_before = Pt(0)
                bullet_p.paragraph_format.space_after = Pt(0)
                bullet_p.paragraph_format.line_spacing = 1.0
                anchor = bullet_p
                bullet_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Spacer between experiences
            spacer = _insert_paragraph_after(anchor, text="", style=None)
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = 1.0
            anchor = spacer
            
        return True

    return False

def _replace_placeholder_bullets(
    doc: DocxDocument,
    placeholder: str,
    lines: list[str],
    bullet_style: str = "Bullet 2",
) -> bool:
    """
    Replace {{PLACEHOLDER}} with multiple bullet paragraphs.

    - The paragraph containing the placeholder becomes the first bullet
    - Remaining bullets are inserted after it
    - Uses an existing bullet style from the template (e.g., "Bullet 2")
    """
    token = f"{{{{{placeholder}}}}}"

    style_name = _first_existing_style(
        doc,
        candidates=[
            bullet_style,
            "Bullet 2",
            "Bullet 1",
            "List Bullet",
            "List Bullet 2",
            "List Paragraph",
            "Viñeta",
            "Lista con viñetas",
        ],
    )

    cleaned = [_strip_leading_bullet_markers(x) for x in lines if (x or "").strip()]
    cleaned = [x for x in cleaned if x]

    for p in _iter_paragraphs(doc):
        if token in p.text:
            _clear_paragraph(p)

            if not cleaned:
                return True
            
            # First bullet
            if style_name:
                p.style = style_name
            p.add_run(cleaned[0])
            _apply_font_to_paragraph(p, FONT_NAME, FONT_SIZE_BODY)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

            # Remaining bullets
            cur = p
            for ln in cleaned[1:]:
                cur = _insert_paragraph_after(cur, ln, style=style_name)
                cur.paragraph_format.space_before = Pt(0)
                cur.paragraph_format.space_after = Pt(0)
                cur.paragraph_format.line_spacing = 1.0

            return True

    return False

# ============================================================
# CV -> blocks
# ============================================================

def _summary_as_paragraph(cv: CVMaster, lang: str) -> str:
    lines = cv.summary.es if lang == "es" else cv.summary.en
    return " ".join(_strip_leading_bullet_markers(line) for line in lines if (line or "").strip()).strip()

def _education_lines(cv: CVMaster, lang: str) -> list[str]:
    out: list[str] = []
    for ed in cv.education:
        degree = ed.degree.es if lang == "es" else ed.degree.en
        location = ed.location.es if lang == "es" else ed.location.en
        out.append(f"{degree} — {ed.institution}, {ed.year}. {location}")
    return out

def _languages_lines(cv: CVMaster, lang: str) -> list[str]:
    out: list[str] = []
    for lng in cv.languages:
        level = lng.level.es if lang == "es" else lng.level.en
        out.append(f"{lng.name} — {level}")
    return out

# ============================================================
# Public API
# ============================================================

TITLES_I18N = {
    "es": {
        "summary_title": "Sobre Mi",
        "Experience_title": "Experiencia",
        "Education_title": "Educación y Certificaciones",
        "skills_title": "Habilidades",
        "language_title": "Idiomas",
    },
    "en": {
        "summary_title": "About Me",
        "Experience_title": "Experience Work",
        "Education_title": "Education & Certifications",
        "skills_title": "Skills",
        "language_title": "Languages",
    },
}


def build_cv_docx(
    cv: CVMaster,
    lang: str,
    out_path: Path,
    skills_ordered: Optional[list[str]] = None,
    bullets_per_experience: Optional[dict[str, list[str]]] = None,
    template_path: str = "templates/cv_template.docx",
) -> None:
    """
    Fills a Word template that contains placeholders like:
      {{NAME}}, {{TITLE}}, {{CONTACT}}, {{SUMMARY}},
      {{EXPERIENCE}}, {{EDUCATION}}, {{SKILLS}}, {{LANGUAGES}}

    IMPORTANT:
    - Ensure your template uses those exact placeholders.
    - Ensure your template has the style "Bullet 2" (Avenir Next Regular 10pt).
    """
    doc: DocxDocument = docx.Document(template_path)
    _set_page_layout(
        doc,
        left_mm=10,
        right_mm=15,
        top_mm=15,
        bottom_mm=15,
        page_w_mm=210,
        page_h_mm=297
        )
    titles = TITLES_I18N[lang]
    contact = cv.profile.contact
    loc = contact.location.es if lang == "es" else contact.location.en
    title = cv.profile.title.es if lang == "es" else cv.profile.title.en
    linkedin_url = contact.links.linkedin.es if lang == "es" else contact.links.linkedin.en
    github_url   = contact.links.github


    _replace_placeholder_name(doc, "NAME", cv.profile.name)
    _replace_placeholder_title(doc, "TITLE", title)
    _replace_placeholder_contact(
      doc,
      "CONTACT",
      email=contact.email,
      location=loc,
      phone=contact.phone,
      linkedin_url=linkedin_url,
      github_url=github_url
  )
    _replace_placeholder_title(doc, "summary_title", titles["summary_title"])
    _replace_placeholder_title(doc, "Experience_title", titles["Experience_title"])
    _replace_placeholder_title(doc, "skills_title", titles["skills_title"])
    _replace_placeholder_title(doc, "language_title", titles["language_title"])

    _replace_placeholder_paragraph(doc, "SUMMARY", _summary_as_paragraph(cv, lang))

    _replace_placeholder_experience_blocks(
        doc,
        "EXPERIENCE",
        cv,
        lang,
        bullets_per_experience,
        job_title_style=None,
        bullet_style_preferred="Bullet 2*",
        font_name=FONT_NAME,
        font_size=FONT_SIZE_BODY,
        job_title_color=COLOR_SUB_TITLE,
    )

    _replace_placeholder_title_hyperlink(
        doc,
        "Education_title",
        text=titles["Education_title"],
        url="https://drive.google.com/drive/u/0/folders/1XmcnXtTeu-2l4snJg5-pUK7orA7iWxBI",
    )
    _replace_placeholder_bullets(doc, "EDUCATION", _education_lines(cv, lang), bullet_style="Bullet 2")

    skills = skills_ordered or (cv.skills.core + cv.skills.apis + cv.skills.tooling)
    _replace_placeholder_inline_left(doc, "SKILLS", ", ".join(skills))

    _replace_placeholder_bullets(doc, "LANGUAGES", _languages_lines(cv, lang), bullet_style="Bullet 2")

    _set_footer_hyperlink_i18n(
    doc,
    text_es="Este CV fue generado por mi propio programa, puedes revisar el codigo fuente aquí",
    text_en="This CV was generated by my own program, you can check the source code here",
    url="https://github.com/FraanMoore?tab=repositories",
    lang=lang,
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

def build_cover_letter_docx(
    cv: CVMaster,
    lang: str,
    company: str,
    role: str,
    out_path: Path,
) -> None:
    doc: DocxDocument = docx.Document()

    name = cv.profile.name
    title = cv.profile.title.es if lang == "es" else cv.profile.title.en
    contact = cv.profile.contact

    doc.add_paragraph(name)
    doc.add_paragraph(title)
    doc.add_paragraph(f"{contact.email} • {contact.phone}")
    doc.add_paragraph("")

    if lang == "es":
        doc.add_paragraph(f"Estimado equipo de {company},")
        doc.add_paragraph("")
        doc.add_paragraph(
            f"Me gustaría postular al cargo de {role}. Soy desarrolladora de software con foco en frontend y experiencia trabajando en aplicaciones web complejas, priorizando mantenibilidad, performance y calidad."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "He participado en el desarrollo y mantención de software, implementación de nuevas funcionalidades, integración con APIs, factorización y optimización, colaborando de forma cercana con equipos multidisciplinarios."
        )
        doc.add_paragraph("")
        doc.add_paragraph("Quedo atenta a coordinar una entrevista. Muchas gracias por su tiempo.")
        doc.add_paragraph("")
        doc.add_paragraph("Saludos,")
        doc.add_paragraph(name)
    else:
        doc.add_paragraph(f"Dear {company} team,")
        doc.add_paragraph("")
        doc.add_paragraph(
            f"I’m interested in applying for the {role} position. I’m a software developer focused on frontend work, experienced in complex web applications, with an emphasis on maintainability, performance, and quality."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "I have worked on software development and maintenance, implemented new features, integrated APIs and business logic, and improved code through refactoring and optimization, collaborating closely with cross-functional teams."
        )
        doc.add_paragraph("")
        doc.add_paragraph("I’d be happy to connect and share more. Thank you for your time and consideration.")
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(name)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
